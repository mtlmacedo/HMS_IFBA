from django.http.response import HttpResponse, FileResponse
from rest_framework import viewsets, permissions, status
from hmsifba.models import *
from hmsifba.serializer import *
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from django.contrib.auth.models import User, Group
from django.contrib.auth import authenticate, login, logout
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
import string
import secrets
from datetime import datetime
from docx import Document

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all().order_by('-date_joined')
    serializer_class = UserSerializer
    permission_classes = [permissions.IsAuthenticated]

class GroupViewSet(viewsets.ModelViewSet):
    queryset = Group.objects.all()
    serializer_class = GroupSerializer
    permission_classes = [permissions.IsAuthenticated]

user_response = openapi.Response('Response Description', UserSerializer)
@api_view(['POST', ])
@permission_classes([AllowAny])
def registration_view(request):

	if request.method == 'POST':
		serializer = RegistrationSerializer(data=request.data)
		data = {}
		if serializer.is_valid():
			cliente = serializer.save()
			data['response'] = 'successfully registered new user.'
			data['email'] = cliente.email
			data['username'] = cliente.username
			token = Token.objects.get(user=cliente).key
			data['token'] = token
		else:
			data = serializer.errors
		return Response(data)

user_response = openapi.Response('Response Description', UserSerializer)
@permission_classes([AllowAny])
def login(request):
    if(request.method == 'POST'):
        serializer = UserSerializer(data=request.data)
        usuario = request.data['username']
        senha = request.data['password']
        user = authenticate(request, username=usuario, password=senha)
        if(user is not None):
            login(request, user)
            return HttpResponse("Success", status=status.HTTP_200_OK)
        else:
            return HttpResponse("Not Found", status=status.HTTP_400_BAD_REQUEST)

@permission_classes([AllowAny])
def logout(resquest):
    try:
        logout(resquest)
        return HttpResponse("Success", status=status.HTTP_200_OK)
    except KeyError:
        return HttpResponse("Not Found", status=status.HTTP_404_NOT_FOUND)

empresa_response = openapi.Response('Response Description', EmpresaSerializer)
@permission_classes([IsAuthenticated])
def get_empresa(request):
    if(request.method == 'GET'):
        empresa = Empresa.objects.all()
        empresa_serializer = EmpresaSerializer(empresa, many=True)
        return Response(empresa_serializer.data)
    elif(request.method == 'POST'):
        empresa_serializer = EmpresaSerializer(data=request.data)
        if(empresa_serializer.is_valid):
            empresa_serializer.save()
            return Response(empresa_serializer.data, status=status.HTTP_201_CREATED)
        return Response(empresa_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@permission_classes([IsAuthenticated])
def detalhar_empresa(request, pk):
    try:
        empresa = Empresa.objects.get(pk=pk)
    except Empresa.DoesNotExist:
        return Response('Empresa não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        empresa_serializer = EmpresaSerializer(empresa)
        return Response(empresa_serializer.data)

    elif request.method == 'PUT':
        empresa_serializer = EmpresaSerializer(empresa, data=request.data)
        if empresa_serializer.is_valid():
            empresa_serializer.save()
            return Response(empresa_serializer.data)
        return Response(empresa_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        empresa.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


quartos_response = openapi.Response('Response Description', QuartoSerializer)
@permission_classes([IsAuthenticated])
def get_quarto(request):
    if request.method == 'GET':
        quarto = Quarto.objects.all()
        quarto_serializer = QuartoSerializer(quarto, many=True)
        return Response(quarto_serializer.data)

    elif request.method == 'POST':
        quarto_serializer = QuartoSerializer(data=request.data)
        if quarto_serializer.is_valid():
            quarto_serializer.save()
            return Response(quarto_serializer.data, status=status.HTTP_201_CREATED)
        return Response(quarto_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@permission_classes([IsAuthenticated])
def detalhar_quarto(request, pk):
    try:
        quarto = Quarto.objects.get(pk=pk)
    except Quarto.DoesNotExist:
        return Response('Quarto não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        quarto_serializer = QuartoSerializer(quarto)
        return Response(quarto_serializer.data)

    elif request.method == 'PUT':
        quarto_serializer = QuartoSerializer(quarto, data=request.data)
        if quarto_serializer.is_valid():
            quarto_serializer.save()
            return Response(quarto_serializer.data)
        return Response(quarto_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    elif request.method == 'DELETE':
        quarto.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

def quartos_disponiveis(capacidade):
    quarto = Quarto.objects.filter(capacidade__icontains = capacidade, disponibilidade = True).first()
    return HttpResponse(quarto.id)

def alterar_status_quarto(pk):
    quarto = Quarto.objects.get(pk=pk)
    if(quarto.disponibilidade is True):
        quarto.disponibilidade = False
        quarto.save()
        return True

colaboradors_response = openapi.Response('Response Description', ColaboradorSerializer)
@permission_classes([IsAuthenticated])
def get_colaborador(request):
    if(request.method == 'GET'):
        # colaborador = Colaborador.objects.all()
        colaborador_serializer = ColaboradorSerializer(Colaborador.objects.all(), many=True)
        return Response(colaborador_serializer.data)
    elif (request.method == 'POST'):
        colaborador_serializer = ColaboradorSerializer(data=request.data)
        if(colaborador_serializer.is_valid()):
            senha = request.data['senha']
            login = request.data['login']
            cargo = request.data['cargo']
            usuario = User.objects.create_user(login, "", senha)

            if(cargo == "G"):
                group = Group.objects.get(name='Gerente')
            elif(cargo == "R"):
                group = Group.objects.get(name='Recepcionista')
            
            group.user_set.add(usuario)
            usuario.is_admin = True
            usuario.is_staff = True
            usuario.save()

            colaborador_serializer.save()
            return Response(colaborador_serializer.data, status=status.HTTP_201_CREATED)
        return Response(colaborador_serializer.data, status=status.HTTP_400_BAD_REQUEST)

@permission_classes([IsAuthenticated])
def detalhar_colaborador(request, pk):
    try:
        colaborador = Colaborador.objects.get(pk=pk)
    except Colaborador.DoesNotExist:
        return Response('colaborador não encontrado', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        colaborador_serializer = ColaboradorSerializer(colaborador)
        return Response(colaborador_serializer.data)

    elif request.method == 'PUT':
        colaborador_serializer = ColaboradorSerializer(colaborador, data=request.data)
        if colaborador_serializer.is_valid():
            colaborador_serializer.save()
            return Response(colaborador_serializer.data)
        return Response(colaborador_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        colaborador.delete()
        return Response(status=status.HTTP_200_OK)

cliente_response = openapi.Response('Response Description', ClienteSerializer)
@permission_classes([IsAuthenticated])
def get_cliente(request):
    if request.method == 'GET':
        cliente = Cliente.objects.all()
        cliente_serializer = ClienteSerializer(cliente, many=True)
        return Response(cliente_serializer.data)

    elif request.method == 'POST':
        cliente_serializer = ClienteSerializer(data=request.data)
        
        if cliente_serializer.is_valid():
            nome = request.data['nome']
            email = request.data['email']
            senha = string.ascii_lowercase + string.digits
            senha = ''.join(secrets.choice(senha) for i in range(8))
            user = User.objects.create_user(nome, email, senha)
            my_group = Group.objects.get(name='Cliente') 
            my_group.user_set.add(user)      
            user.is_admin = True
            user.is_staff = True
            user.save()
            cliente_serializer.save()
            return Response(cliente_serializer.data, status=status.HTTP_201_CREATED)
        return Response(cliente_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@permission_classes([IsAuthenticated])
def detalhar_cliente(request, pk):
    try:
        client = Cliente.objects.get(pk=pk)
    except Cliente.DoesNotExist:
        return Response('Cliente não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        cliente_serializer = ClienteSerializer(client)
        return Response(cliente_serializer.data)

    elif request.method == 'PUT':
        cliente_serializer = ClienteSerializer(client, data=request.data)
        if cliente_serializer.is_valid():
            client.email = request.data.get('email')
            client.endereco = request.data.get('endereco')
            client.telefone = request.data.get('telefone')

            client.save()
            return Response(cliente_serializer.data)
        return Response(cliente_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        client.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

servicos_response = openapi.Response('Response Description', TipoServicoSerializer)
@swagger_auto_schema(method='GET', responses={200: servicos_response})
@swagger_auto_schema(methods=['POST'], request_body=TipoServicoSerializer)
@api_view(['GET', 'POST'])
def get_tipo_servico(request):
    if(request.method == 'GET'):
        tipo_servico = TipoServico.objects.all()
        serializer = TipoServicoSerializer(tipo_servico, many=True)
        return Response(serializer.data)
    elif (request.method == 'POST'):
        serializer = TipoServicoSerializer(data=request.data)
        if(serializer.is_valid()):
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@permission_classes([IsAuthenticated])
def detalhar_tipo_servico(request, pk):
    try:
        tipo_servico = TipoServico.objects.get(pk=pk)
    except TipoServico.DoesNotExist:
        return Response('Tipo de Serviço não encontrado', status=status.HTTP_404_NOT_FOUND)
    if (request.method == 'GET'):
        serializer = TipoServicoSerializer(tipo_servico)
        return Response(serializer.data)
    elif (request.method == 'PUT'):
        serializer = TipoServicoSerializer(tipo_servico, data=request.data)
        if(serializer.is_valid()):
            serializer.save()
            return Response(serializer.data)
        return Response(status=status.HTTP_400_BAD_REQUEST)
    elif (request.method == 'DELETE'):
        tipo_servico.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

reservas_response = openapi.Response('Response Description', ReservaSerializer)
@permission_classes([IsAuthenticated])
def get_reserva(request):
    if (request.method == 'GET'):
        reserva = Reserva.objects.all()
        serializer = ReservaSerializer(reserva, many=True)
        return Response(serializer.data)
    elif (request.method == 'POST'):
        serializer = ReservaSerializer(data=request.data)
        if (serializer.is_valid()):
            id_quarto = request.data['quarto']
            quarto_reservado = alterar_status_quarto(id_quarto)
            if(quarto_reservado):
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            else:
                return Response("Quarto está ocupado", status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
@permission_classes([IsAuthenticated])
def detalhar_reserva(request, pk):
    try:
        reserva = Reserva.objects.get(pk=pk)
    except Reserva.DoesNotExist:
        return Response("Reserva não encontrada", status=status.HTTP_404_NOT_FOUND)
    if (request.method == 'GET'):
        serializer = ReservaSerializer(reserva)
        return Response(serializer.data)
    elif (request.method == 'PUT'):
        serializer = ReservaSerializer(reserva, data=request.data)
        if(serializer.is_valid()):
            id_quarto = request.data['quarto']
            total_pessoas = request.data['qtd_pessoas']
            if(total_pessoas > reserva.qtd_pessoas and id_quarto == reserva.quarto.id):
                return Response("Quantidade de pessoas não suportada", status=status.HTTP_400_BAD_REQUEST)
            if(id_quarto != reserva.quarto.id):
                nova_reserva = alterar_status_quarto(id_quarto)
                if(nova_reserva):
                    quarto = reserva.quarto
                    quarto.disponibilidade
                    quarto.save()
                    return Response(serializer.data)
                else:
                    serializer.save()
                    return Response("Quarto ocupado", status=status.HTTP_400_BAD_REQUEST)
            return Response(status=status.HTTP_400_BAD_REQUEST)
    elif (request.method == 'DELETE'):
        reserva.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

def criarReserva(request):
    reserva = Reserva()
    reserva.cliente = Cliente.objects.get(pk=request.data['cliente']) 
    reserva.dataEntrada = request.data['dataEntrada']
    reserva.dataSaida = request.data['dataSaida']
    reserva.qtd_pessoas = request.data['qtd_pessoas']
    return reserva

estadias_response = openapi.Response('Response Description', EstadiaSerializer)
@permission_classes([IsAuthenticated])
def get_estadia(request):
    if(request.method == 'GET'):
        estadia = Estadia.objects.all()
        serializer = EstadiaSerializer(estadia, many=True)
        return Response(serializer.data)
    elif (request.method == 'POST'):
        serializer = EstadiaSerializer(data=request.data)
        reserva_id = request.data['reserva']        
        if (reserva_id is None):
            estadia = criarEstadia(request)
            reserva = criarReserva(request)
            quarto_id = quartos_disponiveis(request.data['quantidade_pessoas'])
            if (quarto_id is not None):
                reservado = alterar_status_quarto(quarto_id)
                estadia = atualizarEstadia(reservado, reserva, estadia, quarto_id)
                estadia = Estadia.objects.get(pk=estadia.id)
                serializer = EstadiaSerializer(estadia)
                return Response(serializer.data)
        else:
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@permission_classes([IsAuthenticated])
def detalhar_estadia(request, pk):
    try:
        estadia = Estadia.objects.get(pk=pk)
    except Estadia.DoesNotExist:
        return Response('Estadia não encontrada', status=status.HTTP_404_NOT_FOUND)
    if request.method == 'GET':
        serializer = EstadiaSerializer(estadia)
        return Response(serializer.data)
    elif request.method == 'PUT':
        serializer = EstadiaSerializer(estadia, data=request.data)
        if serializer.is_valid():
            if estadia.isMudancaDeQuarto is True:
                quarto = mudar_quarto(estadia, pk) 
                if quarto is not None:
                    estadia.isMudancaDeQuarto = False
                    estadia.save()
                    return Response(serializer.data)
                else:
                    return Response("Não há quartos que satisfaçam a solicitação",status=status.HTTP_404_NOT_FOUND)
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

def criarEstadia(request):
    estadia = Estadia()
    estadia.numero_cartao = request.data['numero_cartao']
    estadia.dataEntrada = request.data['dataEntrada']
    estadia.dataSaida = request.data['dataSaida']
    estadia.qtd_pessoas = request.data['qtd_pessoas']
    estadia.qtd_quartos = request.daa['qtd_quartos']
    estadia.servico = request.data['servico']
    return estadia

def atualizarEstadia(reservado, reserva, estadia, quarto_id):
     if reservado is True:
        quarto = Quarto.objects.get(pk=quarto_id)
        reserva.quarto = quarto
        reserva.save()
        estadia.cliente = reserva.cliente
        estadia.reserva = reserva
        estadia.save()
        estadia = Estadia.objects.get(pk=estadia.id)
        return estadia

def gerar_fatura(valor_total, pk):
    estadia = Estadia.objects.get(pk=pk)

    document = Document()
    docx_title=f"fatura{estadia.cliente.nome}.docx"

    document.add_heading('Fatura', 0)

    title = document.add_paragraph('Fatura referente à estadia no Hotel IFBA ')

    document.add_heading('Dados da Estadia', 1)
    document.add_paragraph(f'Data de Entrada: {estadia.entrada}')
    document.add_paragraph(f'Data de Saída: {estadia.checkout}')
    document.add_paragraph(f'Serviço Prestado: {estadia.servico.tipo } ')
    document.add_paragraph(f'Valor diário do serviço: {estadia.servico.preco } ')
    document.add_paragraph(f'Quantidade de Pessoas: {estadia.quantidade_pessoas } ')

    document.add_heading('Dados do Cliente',1)
    document.add_paragraph(f'Nome do cliente: {estadia.cliente.nome}')
    document.add_paragraph(f'Número do cartão: {estadia.cartao}')
    document.add_paragraph(f'Telefone: {estadia.cliente.telefone}')

    document.add_heading('Valor a pagar',1)
    document.add_paragraph(f'R$: {valor_total}')

    document.save(f'hmsifba/static/media/{docx_title}')      
    return document

def baixar_fatura(request, pk):
     if request.method == 'GET':
        estadia = Estadia.objects.get(pk=pk)
        nome_cliente = estadia.cliente.nome
        nome_doc = f'fatura{nome_cliente}.docx'
        file_path =  f'hmsifba/static/media/{nome_doc}'
        file = open(file_path, 'rb') 
        response = FileResponse(file)
        response['Content-Type'] = 'application/application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response['Content-Disposition'] = f'attachment;filename={nome_doc}'      

        return response

def mudar_quarto(estadia, pk):
    quarto = Quarto.objects.filter(capacidade__icontains = estadia.quantidade_pessoas, isDisponivel = True).first()
    quarto_atual = estadia.reserva.quarto

    if quarto is not None:
        reserva_quarto = alterar_status_quarto(quarto.id)
        
        if reserva_quarto is True:
            quarto_atual.isDisponivel = True
            quarto_atual.save()
            estadia.reserva.quarto = quarto          
            estadia.reserva.save()

    return quarto

estatistica_response = openapi.Response('Response Description', EstatisticaSerializer)
@permission_classes([IsAuthenticated])
def get_estatistica(request):
    if request.method == 'GET':
        estatistica = Estatistica.objects.all()
        serializer = EstatisticaSerializer(estatistica, many=True)
        return Response(serializer.data)
        
    return Response(status=status.HTTP_400_BAD_REQUEST)

def createEstatistica(valor_total, entrada, cliente):
    data_entrada = datetime.strptime(entrada.strftime('%m/%d/%Y'), "%m/%d/%Y")
    anoReferente = data_entrada.year
    mes = data_entrada.month
    semestreReferente = verificarSemestre(mes)
    estatistica = Estatistica.objects.filter(semestre=semestreReferente, ano=anoReferente)
 
    if (estatistica.exists()):
        estatistica = Estatistica.objects.get(semestre=semestreReferente, ano=anoReferente)
        estatistica.ano = anoReferente
        estatistica.semestre = semestreReferente
        estatistica.faturamentoAnual = faturamentoAnual(valor_total)
        estatistica.faturamentoDoTrimestre = faturamentoSemestre(valor_total, semestreReferente)
        estatistica.taxaQuartosVendidos = taxaQuartosVendidos()
        estatistica.custoTotalCliente = obterCustoTotalCliente(cliente, valor_total)
        estatistica.clientePremium = obter_cliente(cliente)
        estatistica.clienteId = cliente.id
        estatistica.save()
    else:
        estatistica = Estatistica()
        estatistica.ano = anoReferente
        estatistica.semestre = semestreReferente
        estatistica.faturamentoAnual = faturamentoAnual(valor_total)
        estatistica.faturamentoDoTrimestre = faturamentoSemestre(valor_total, semestreReferente)
        estatistica.taxaQuartosVendidos = taxaQuartosVendidos()
        estatistica.clienteId = cliente.id
        estatistica.custoTotalCliente = valor_total
        estatistica.clientePremium = obter_cliente(cliente)
        estatistica.save()

def verificarSemestre(mes):
    if(mes >= 1 and mes <=6):
        return 'primeiro semestre'
    elif (mes >=7 and mes <= 12):
        return 'segundo semestre'
    return

def faturamentoSemestre(valor_total, semestreReferente):
    faturamentoSemestre = Estatistica.objects.filter(semestre = semestreReferente).aggregate(sum('faturamentoSemestre'))['faturamentoSemestre__sum']

    if faturamentoSemestre is None:
        return valor_total
    else:
        return faturamentoSemestre + float(valor_total)

def faturamentoAnual(valor_total):
    fatAnual = Estatistica.objects.all().aggregate(sum('faturamentoAnual'))['faturamentoAnual__sum']

    if fatAnual is None:
        return valor_total
    else:
        return fatAnual + float(valor_total)

def taxaQuartosVendidos():
    quartoVendidos = Quarto.objects.filter(disponibilidade = False).count()*100
    quartoTotal = Quarto.objects.all().count()

    return quartoVendidos/quartoTotal

def obter_cliente(cliente):
    estatistica_cliente = Estatistica.objects.all().order_by('custoTotalCliente').first()

    if estatistica_cliente is None:
        return cliente.nome
    else:
        cliente = Cliente.objects.get(pk = estatistica_cliente.id)
        return cliente.nome

def obterCustoTotalCliente(cliente, valor_total):
    estatistica_cliente = Estatistica.objects.filter(clienteId = cliente.id).order_by('custoTotalCliente').first()

    if estatistica_cliente is None:
        return valor_total
    else:
        estatistica_cliente.custoTotalCliente = float(valor_total) + estatistica_cliente.custoTotalCliente
        return estatistica_cliente.custoTotalCliente